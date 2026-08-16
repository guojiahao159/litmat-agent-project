"""生成完善后的初赛方案docx（基于模板结构 + 项目实际实现）"""

import sys

sys.stdout.reconfigure(encoding="utf-8")

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = docx.Document()

# ========== 标题 ==========
title = doc.add_heading("材料科学文献调研Agent_算法赛初赛方案", level=0)
subtitle = doc.add_paragraph("——基于多Agent协作的固态电解质构效关系发现系统")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== 一、项目概述 ==========
doc.add_heading("一、项目概述", level=1)
doc.add_heading("1.1 项目名称", level=2)
doc.add_paragraph("LitMat-Agent：面向固态电解质研发的文献驱动科学发现智能体系统")

doc.add_heading("1.2 参赛方向", level=2)
doc.add_paragraph("方向三：材料科学文献驱动的科学发现智能体（基本任务：文献调研Agent + 路线A：构效关系发现）")

doc.add_heading("1.3 方案概述", level=2)
doc.add_paragraph(
    "本项目针对固态电解质领域知识分散、构效关系复杂、研发周期长的痛点，设计并实现一个基于"
    "\"工作流编排+多Agent混合\"架构的文献驱动科学发现系统。系统以Sci-Base本地文献库和Sciverse API为数据源，"
    "通过任务规划、文献检索、知识抽取、跨文献融合、Research Gap识别、证据核验等9个专业化Agent协作，"
    "实现从科学问题输入到可验证构效关系假设输出的全流程自动化。项目聚焦固态电解质（硫化物/氧化物/聚合物）"
    "的离子电导率、界面稳定性等关键性能，旨在建立\"成分-结构-工艺-性能\"的定量构效关系，"
    "为新型固态电解质设计提供数据驱动的决策支持。"
)
doc.add_paragraph(
    "当前已完成可运行原型：9个Agent全部独立实现并接入LangGraph工作流，"
    "形成\"规划→检索→筛选→解析→抽取→分析→核验→报告\"8节点端到端链路，"
    "配套Streamlit交互界面与证据链追踪机制，支持未配置LLM API时的演示模式运行。"
)

# ========== 二、科学问题理解 ==========
doc.add_heading("二、科学问题理解", level=1)
doc.add_heading("2.1 科学问题与研究对象", level=2)
doc.add_paragraph(
    "固态电解质是下一代高安全、高能量密度全固态电池的核心材料，其研发面临三大科学挑战："
    "（1）离子输运机制复杂：锂离子在晶格中的迁移受晶体结构、缺陷浓度、晶界效应等多因素耦合影响，"
    "现有理论模型难以准确预测；（2）构效关系碎片化：文献中报道的\"成分-结构-性能\"关联分散于数千篇论文，"
    "缺乏系统整合，且存在大量矛盾结论（如LLZO的Al掺杂vs Ta掺杂效果争议）；"
    "（3）知识发现效率低：传统试错法研发周期长达10-20年，而文献调研依赖人工阅读，覆盖率不足10%，"
    "大量潜在关联未被挖掘。"
)
doc.add_paragraph(
    "本项目的核心科学问题是：如何从大规模材料科学文献中自动抽取、融合、推理固态电解质的构效关系，"
    "识别具有新颖性和可证伪性的Research Gap，加速\"文献知识→科学假设→实验验证\"的闭环。"
    "研究对象聚焦三类主流固态电解质：硫化物（如Li10GeP2S12、Li3PS4）、氧化物（如LLZO、LATP）"
    "和聚合物（如PEO基复合电解质）。"
)

doc.add_heading("2.2 科学意义", level=2)
doc.add_paragraph(
    "本项目的科学意义体现在三个层面：（1）方法论层面：建立文献驱动科学发现的Agent范式，"
    "将大语言模型的推理能力、信息抽取技术与材料领域知识深度融合，推动AI for Science从\"辅助工具\"向"
    "\"自主发现\"演进；（2）领域层面：构建固态电解质领域首个大规模构效关系知识库，"
    "揭示隐藏的成分-结构-性能关联，为理性设计提供理论指导；（3）应用层面：缩短新型固态电解质研发周期，"
    "降低研发成本，支撑我国全固态电池技术的自主创新与产业化进程。"
)
doc.add_paragraph(
    "与现有系统相比，本项目的差异化优势在于：普通文献检索仅返回相关论文列表，"
    "本系统提供结构化知识图谱与可验证假设；普通RAG系统基于单篇或少量文献生成答案，"
    "本系统实现跨文献知识融合与矛盾检测；通用文献问答系统缺乏领域深度，"
    "本系统内置材料科学本体与专家规则，确保抽取的化学成分、晶体结构、性能数据等专业信息的准确性。"
)

# ========== 三、技术方案与预期方法路线 ==========
doc.add_heading("三、技术方案与预期方法路线", level=1)
doc.add_heading("3.1 技术方案", level=2)
doc.add_paragraph(
    "系统采用\"工作流编排+多Agent混合\"架构，兼顾流程可控性与任务灵活性。"
    "整体架构分为四层：数据层、Agent层、工作流层、应用层。"
)

# 表格1：系统架构组件与功能
table1 = doc.add_table(rows=1, cols=4)
table1.style = "Table Grid"
hdr = table1.rows[0].cells
hdr[0].text = "层级"
hdr[1].text = "组件"
hdr[2].text = "功能"
hdr[3].text = "技术选型"
rows1 = [
    ("数据层", "Sci-Base本地库", "2500万+篇文献全文存储与检索", "Elasticsearch"),
    ("数据层", "Sciverse API", "实时语义检索与全文定位", "RESTful API"),
    ("数据层", "结构化知识库", "材料实体、属性、关系统一存储", "PostgreSQL + Neo4j"),
    ("Agent层", "任务规划Agent", "科学问题理解与任务拆解", "GPT-4/Claude + Prompt工程"),
    ("Agent层", "文献检索Agent", "多策略检索与结果聚合", "Embedding + Reranker"),
    ("Agent层", "文献筛选Agent", "相关性评估与去重", "关键词加权 + DOI去重"),
    ("Agent层", "PDF解析Agent", "全文结构化解析", "MinerU + 自定义规则"),
    ("Agent层", "知识抽取Agent", "成分/结构/性能/工艺抽取", "LLM + 材料本体 + 规则"),
    ("Agent层", "知识融合Agent", "跨文献实体对齐与融合", "实体链接 + 图融合算法"),
    ("Agent层", "Gap识别Agent", "Research Gap检测与评估", "对比分析 + 新颖性评估"),
    ("Agent层", "证据核验Agent", "文献溯源与事实核验", "引用链追踪 + 交叉验证"),
    ("Agent层", "报告生成Agent", "结构化调研报告生成", "模板引擎 + LLM润色"),
    ("工作流层", "流程编排引擎", "任务调度与状态管理", "LangGraph"),
    ("应用层", "交互界面", "用户查询与结果展示", "Streamlit"),
]
for r in rows1:
    row = table1.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v
doc.add_paragraph("表1 系统架构组件与功能").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("3.2 预期方法路线", level=2)
doc.add_paragraph(
    "系统运行流程分为16个步骤，形成完整的\"输入-处理-输出\"闭环。当前已实现8节点核心链路："
    "任务规划→文献检索→文献筛选→PDF解析→知识抽取→分析→证据核验→报告生成。"
)

# 表格2：系统运行流程与输入输出
table2 = doc.add_table(rows=1, cols=5)
table2.style = "Table Grid"
hdr = table2.rows[0].cells
hdr[0].text = "步骤"
hdr[1].text = "处理环节"
hdr[2].text = "输入"
hdr[3].text = "处理方法"
hdr[4].text = "输出"
rows2 = [
    ("1", "科学问题输入", "用户自然语言查询", "意图识别与实体抽取", "结构化查询表示"),
    ("2", "任务拆解", "结构化查询", "LLM任务规划", "子任务清单"),
    ("3", "检索策略生成", "子任务描述", "查询扩展与同义词映射", "多路检索查询"),
    ("4", "多源文献检索", "检索查询", "Sci-Base本地检索+Sciverse API调用", "候选文献集"),
    ("5", "去重与筛选", "候选文献", "标题/摘要/全文三级筛选", "相关文献集"),
    ("6", "PDF全文解析", "PDF文件", "MinerU解析+版面分析", "结构化文本"),
    ("7", "材料知识抽取", "结构化文本", "LLM+规则混合抽取", "材料实体与属性"),
    ("8", "实体规范化", "原始实体", "材料本体映射+单位统一", "标准化知识单元"),
    ("9", "知识库存储", "标准化知识", "图数据库写入", "材料知识图谱"),
    ("10", "跨文献融合", "多篇文献知识", "实体对齐+冲突检测", "融合知识视图"),
    ("11", "冲突与缺失检测", "融合知识", "对比分析+统计检测", "矛盾点/缺失点清单"),
    ("12", "Research Gap生成", "矛盾/缺失清单", "Gap模式匹配+新颖性评估", "候选Gap列表"),
    ("13", "文献证据核验", "候选Gap", "引用链追踪+原文定位", "证据链报告"),
    ("14", "Gap评分排序", "核验后Gap", "多维度评分模型", "优先级排序"),
    ("15", "调研报告生成", "排序后Gap+证据", "模板填充+LLM润色", "结构化报告"),
    ("16", "引用与事实检查", "生成报告", "交叉引用验证", "最终调研报告"),
]
for r in rows2:
    row = table2.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v
doc.add_paragraph("表2 系统运行流程与输入输出").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "各组件职责划分：大语言模型（GPT-4/Claude）负责任务规划、知识抽取、Gap推理与报告生成；"
    "Embedding模型（BGE-M3）负责语义检索与相似度计算；Reranker（BGE-Reranker）负责检索结果精排；"
    "MinerU负责PDF深度解析；PostgreSQL存储结构化知识，Neo4j存储知识图谱；"
    "规则程序负责单位统一、化学式规范化等确定性任务；LangGraph负责工作流编排与状态管理。"
)

doc.add_heading("3.3 数据来源、依赖工具与运行流程", level=2)
doc.add_heading("（1）数据来源", level=3)
doc.add_paragraph(
    "本项目采用\"Sci-Base本地+Sciverse API\"混合数据策略：Sci-Base（opendatalab/Sci-Base）"
    "提供2500万+篇开放获取文献的本地存储，保障数据隐私与处理效率；Sciverse API提供4.66亿条学术元数据"
    "与2800万+全文证据片段，支持实时语义检索与全文定位，其调用记录天然构成可审计的证据链。"
    "两者结合，兼顾数据规模、检索质量与合规要求。"
)

doc.add_heading("（2）依赖工具", level=3)
# 表格3：核心依赖工具与许可证
table3 = doc.add_table(rows=1, cols=5)
table3.style = "Table Grid"
hdr = table3.rows[0].cells
hdr[0].text = "工具"
hdr[1].text = "用途"
hdr[2].text = "开源情况"
hdr[3].text = "许可证"
hdr[4].text = "复现影响"
rows3 = [
    ("Sci-Base", "本地文献库", "开放数据", "CC BY 4.0", "可自由使用"),
    ("Sciverse API", "语义检索", "商业API", "服务协议", "需注册账号"),
    ("MinerU", "PDF解析", "开源", "MIT", "可自由使用"),
    ("BGE-M3", "Embedding", "开源", "MIT", "可自由使用"),
    ("BGE-Reranker", "结果精排", "开源", "MIT", "可自由使用"),
    ("GPT-4/Claude", "LLM推理", "商业API", "服务协议", "需API Key"),
    ("PostgreSQL", "关系数据库", "开源", "PostgreSQL License", "可自由使用"),
    ("Neo4j", "图数据库", "社区版开源", "GPLv3", "社区版可自由使用"),
    ("LangGraph", "工作流引擎", "开源", "MIT", "可自由使用"),
    ("Docker", "容器化部署", "开源", "Apache 2.0", "可自由使用"),
]
for r in rows3:
    row = table3.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v
doc.add_paragraph("表3 核心依赖工具与许可证").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("（3）运行流程", level=3)
doc.add_paragraph(
    "系统部署采用Docker Compose编排，包含以下服务：（1）数据导入服务：从Hugging Face下载Sci-Base数据，"
    "经MinerU解析后写入Elasticsearch；（2）Agent服务：各专业化Agent以微服务形式运行，通过gRPC通信；"
    "（3）工作流服务：LangGraph引擎调度Agent执行顺序，处理异常与重试；（4）API网关：对外提供RESTful接口，"
    "接收用户查询并返回调研报告。典型查询响应时间目标：简单查询<30秒，复杂调研任务<5分钟。"
)

# ========== 四、阶段性实验结果或可行性验证 ==========
doc.add_heading("四、阶段性实验结果或可行性验证", level=1)
doc.add_heading("4.1 阶段性实验或可行性验证", level=2)
doc.add_paragraph(
    "本项目已完成可运行原型构建，以下为可行性验证方案与初步验证结果。"
)
doc.add_heading("（1）验证目标", level=3)
doc.add_paragraph(
    "验证系统在固态电解质细分领域的端到端能力，重点评估：文献检索的覆盖率与准确率、"
    "知识抽取的精度、Research Gap识别的新颖性与可操作性。"
)
doc.add_heading("（2）验证设计", level=3)
doc.add_paragraph(
    "选择\"硫化物固态电解质的离子电导率优化\"作为验证场景，该领域文献量适中（约500-800篇核心论文）、"
    "构效关系明确、有公开基准数据集。设计对比实验："
)
# 表格4：可行性验证实验设计
table4 = doc.add_table(rows=1, cols=4)
table4.style = "Table Grid"
hdr = table4.rows[0].cells
hdr[0].text = "实验组"
hdr[1].text = "方法"
hdr[2].text = "评估指标"
hdr[3].text = "预期对比"
rows4 = [
    ("Baseline-1", "关键词检索（BM25）", "Precision/Recall/nDCG@10", "基准值"),
    ("Baseline-2", "语义检索（BGE-M3）", "Precision/Recall/nDCG@10", "验证语义检索优势"),
    ("Baseline-3", "普通RAG（单篇文献问答）", "答案准确率、引用完整性", "验证跨文献融合必要性"),
    ("Ours-1", "混合检索（关键词+语义+Rerank）", "Precision/Recall/nDCG@10", "验证混合策略优势"),
    ("Ours-2", "单Agent端到端", "任务完成率、报告质量", "验证多Agent必要性"),
    ("Ours-3", "完整多Agent系统", "Gap新颖性、专家认可度", "核心验证目标"),
]
for r in rows4:
    row = table4.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v
doc.add_paragraph("表4 可行性验证实验设计").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("（3）评估指标", level=3)
doc.add_paragraph(
    "检索质量：Precision@K、Recall@K、nDCG@K（K=5,10,20）；知识抽取：实体识别F1、关系抽取F1、"
    "数值与单位匹配准确率；Gap质量：专家盲评新颖性（1-5分）、可操作性（1-5分）、证据完整率；"
    "系统效率：端到端响应时间、API调用成本。"
)
doc.add_heading("（4）数据准备", level=3)
doc.add_paragraph(
    "从Sci-Base筛选硫化物固态电解质相关文献500篇，人工标注100篇作为测试集"
    "（包含材料成分、电导率、活化能等关键信息），邀请2-3名材料领域研究生作为专家评估Gap质量。"
)

doc.add_heading("4.2 当前结果", level=2)
doc.add_paragraph(
    "当前已完成系统原型构建与端到端链路验证，主要进展如下："
)
doc.add_paragraph(
    "（1）系统架构落地：9个专业化Agent全部独立实现（任务规划、文献检索、文献筛选、PDF解析、"
    "知识抽取、知识融合、Gap识别、证据核验、报告生成），并接入LangGraph工作流形成8节点端到端链路"
    "（规划→检索→筛选→解析→抽取→分析→核验→报告）。", style="List Bullet"
)
doc.add_paragraph(
    "（2）多源检索实现：Sci-Base本地Elasticsearch检索与Sciverse API实时检索联动，"
    "本地检索失败或结果不足时自动调用Sciverse补充，按标题去重合并，实现方案承诺的多源检索策略。", style="List Bullet"
)
doc.add_paragraph(
    "（3）知识抽取与单位统一：基于正则与规则的材料化学式识别、性能数据抽取，"
    "电导率单位自动统一为S/cm（如1.2 mS/cm → 0.0012 S/cm），支持中英文双语查询。", style="List Bullet"
)
doc.add_paragraph(
    "（4）证据链追踪：为每个Research Gap与候选假设构建可审计证据链，"
    "支持引用原文定位与核验，报告自动生成证据链字段，满足方案\"调用记录天然构成可审计证据链\"的要求。", style="List Bullet"
)
doc.add_paragraph(
    "（5）交互界面：Streamlit Web界面支持研究问题输入、参数设置、示例查询，"
    "调研报告以Markdown渲染展示，统计指标卡片化呈现，支持未配置LLM API时的演示模式运行。", style="List Bullet"
)
doc.add_paragraph(
    "（6）数据持久化：PostgreSQL存储文献/材料/性能结构化数据，Neo4j构建材料知识图谱，"
    "失败时优雅降级不影响主流程。", style="List Bullet"
)
doc.add_paragraph(
    "后续计划：完成对比实验（混合检索vs单一检索的Recall提升、多Agent vs单Agent的Gap质量提升），"
    "补充系统端到端响应时间基线与运行截图。"
)

# ========== 五、复现与开放计划 ==========
doc.add_heading("五、复现与开放计划", level=1)
doc.add_heading("5.1 复现方式", level=2)
doc.add_paragraph(
    "本项目提供完整的复现方案：（1）代码仓库：托管于GitHub，包含完整源代码、依赖清单（pyproject.toml）"
    "与环境配置模板；（2）数据准备：提供Sci-Base数据下载脚本与预处理流程，支持小规模测试数据集快速验证；"
    "（3）一键启动：执行uv sync安装依赖后，通过uv run python main.py启动主程序，"
    "或uv run streamlit run app.py启动Web交互界面；（4）文档说明：提供详细的部署文档、API文档与使用示例。"
    "硬件要求：最低8GB内存+20GB磁盘，推荐16GB内存+GPU（用于本地LLM推理）。"
)

doc.add_heading("5.2 开源计划", level=2)
doc.add_paragraph(
    "本项目采用\"核心代码开源+数据合规使用\"的开放策略：（1）开源范围：Agent框架、工作流引擎、"
    "知识抽取规则、评估脚本全部开源，采用MIT许可证；（2）数据边界：Sci-Base数据遵循CC BY 4.0协议，"
    "用户需自行从Hugging Face下载；Sciverse API需用户自行注册获取访问权限；（3）模型边界："
    "默认支持开源LLM（如Qwen2.5-72B、Llama3.1-70B），商业API（GPT-4/Claude）为可选配置，"
    "用户需自行申请API Key；（4）社区建设：项目主页提供文档、教程、示例，"
    "欢迎社区贡献材料领域本体与抽取规则。"
)

doc.add_heading("5.3 依赖、数据来源与合规披露", level=2)
doc.add_paragraph(
    "（1）开源依赖：本项目基于Python 3.14开发，核心依赖包括LangGraph（MIT）、MinerU（MIT）、"
    "Transformers（Apache 2.0）、Streamlit（Apache 2.0）等，完整依赖清单见pyproject.toml，"
    "所有依赖均允许商业使用。"
)
doc.add_paragraph(
    "（2）数据来源合规：Sci-Base数据集（opendatalab/Sci-Base）采用CC BY 4.0许可证，"
    "允许自由使用与修改，需署名原始来源；Sciverse API为商业服务，使用需遵守服务协议，"
    "调用记录将完整保存以满足可审计性要求。"
)
doc.add_paragraph(
    "（3）商业API使用范围：本项目默认配置使用开源LLM（Qwen2.5-72B），商业API（GPT-4/Claude）"
    "仅作为可选增强配置。若使用商业API，将在文档中明确标注使用范围（仅用于知识抽取与Gap推理），"
    "并提供开源替代方案的部署指南，确保无商业API时系统仍可运行（性能略有下降）。"
)
doc.add_paragraph(
    "（4）AI辅助声明：本项目方案撰写过程中使用了AI工具进行文献调研与文本润色，"
    "但核心架构设计、技术选型、科学问题分析由团队独立完成。赛事期间的所有代码实现与实验验证"
    "将由团队独立完成，符合赛事关于AI辅助的规定。"
)

# ========== 六、团队介绍 ==========
doc.add_heading("六、团队介绍", level=1)
doc.add_heading("6.1 成员背景", level=2)
doc.add_paragraph("【待补充】团队成员姓名、学校/公司、专业背景、核心技能。")
doc.add_heading("6.2 团队分工", level=2)
doc.add_paragraph("【待补充】各成员具体分工与职责。")
doc.add_heading("6.3 团队成果", level=2)
doc.add_paragraph("【待补充】团队过往项目经历与获奖情况。")

# ========== 附录 ==========
doc.add_heading("附录：待团队补充信息清单", level=1)
table5 = doc.add_table(rows=1, cols=4)
table5.style = "Table Grid"
hdr = table5.rows[0].cells
hdr[0].text = "序号"
hdr[1].text = "待补充内容"
hdr[2].text = "所属章节"
hdr[3].text = "优先级"
rows5 = [
    ("1", "团队成员姓名、学校/公司、专业背景", "6.1 成员背景", "高"),
    ("2", "各成员具体分工与职责", "6.2 团队分工", "高"),
    ("3", "团队过往项目经历与获奖情况", "6.3 团队成果", "高"),
    ("4", "可行性验证实验的具体结果数据", "4.2 当前结果", "高（复赛阶段）"),
    ("5", "系统原型的实际运行截图或演示视频", "4.2 当前结果", "中（复赛阶段）"),
    ("6", "与领域专家的合作意向或顾问支持", "6.1 成员背景", "中"),
]
for r in rows5:
    row = table5.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

output_path = r"c:\Users\王一博\Desktop\datawhale02\datawhale_ceshi01\材料科学文献调研Agent_算法赛初赛方案_完善版.docx"
doc.save(output_path)
print(f"已生成: {output_path}")
