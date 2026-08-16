"""提取方案docx的段落与表格内容"""

import sys

sys.stdout.reconfigure(encoding="utf-8")

import docx

doc = docx.Document(
    r"c:\Users\王一博\Desktop\datawhale02\datawhale_ceshi01\材料科学文献调研Agent_算法赛初赛方案.docx"
)

print("=== 段落 ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(f"[{p.style.name}] {p.text}")

print()
print(f"=== 表格数: {len(doc.tables)} ===")
for ti, table in enumerate(doc.tables, 1):
    print(f"\n--- 表格 {ti} ---")
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        print(" | ".join(cells))
