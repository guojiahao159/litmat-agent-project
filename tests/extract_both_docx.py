"""提取两个docx的内容：方案文档 + 模板文档"""

import sys

sys.stdout.reconfigure(encoding="utf-8")

import docx

def extract(path, label):
    print(f"\n{'='*60}")
    print(f"=== {label} ===")
    print(f"{'='*60}")
    doc = docx.Document(path)
    print("--- 段落 ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"[{p.style.name}] {p.text}")
    print(f"\n--- 表格数: {len(doc.tables)} ---")
    for ti, table in enumerate(doc.tables, 1):
        print(f"\n表格 {ti}:")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            print(" | ".join(cells))

extract(
    r"c:\Users\王一博\Desktop\datawhale02\datawhale_ceshi01\材料科学文献调研Agent_算法赛初赛方案.docx",
    "方案文档",
)
extract(
    r"c:\Users\王一博\Downloads\AI for reserach算法赛初赛模板.docx",
    "模板文档",
)
